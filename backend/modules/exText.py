import pymupdf
import re
import os
from docx import Document
import mammoth

from modules.multi_column import column_boxes  # Existing function

def extract_text_from_pdf(file_path):
    """Extracts text from a PDF using PyMuPDF."""
    text = ''
    doc = pymupdf.open(file_path)
    for page in doc:
        bboxes = column_boxes(page, footer_margin=50, no_image_text=True)
        for rect in bboxes:
            text += page.get_text(clip=rect, sort=True) + " "

    # Replace '|' with space
    text = re.sub(r'\|', ' ', text)

    # Allow '@', single hyphen '-', en dash '–', and plus sign '+'
    text = re.sub(r'[^a-zA-Z0-9/.,()@: \n\-–+]', '', text)  # Clean unwanted characters
    text = re.sub(r'\s+', ' ', text)  # Normalize spaces
    print("-------------------------------")
    print(f"RESUME TEXT: {text}")

    return text.strip()

# import pdfplumber

# def extract_text_from_pdf(path):
#     text = ""
#     with pdfplumber.open(path) as pdf:
#         for page in pdf.pages:
#             text += page.extract_text()
#     return text


# def extract_text_from_docx(file_path):
#     """Extracts text from a DOCX file using python-docx."""
#     doc = docx.Document(file_path)
#     text = '\n'.join([para.text for para in doc.paragraphs])
#     print("-------------------------------")
#     print(f"RESUME TEXT: {text}")
    
#     return text.strip()

# def extract_text_from_docx(file_path):
#     """
#     Extracts all visible text from a DOCX file, including paragraphs and table content.
#     Removes duplicate consecutive lines for cleaner output.
#     """
#     try:
#         doc = Document(file_path)
#         text_parts = []

#         # Extract paragraphs
#         for para in doc.paragraphs:
#             if para.text.strip():
#                 text_parts.append(para.text.strip())

#         # Extract text from tables
#         for table in doc.tables:
#             for row in table.rows:
#                 for cell in row.cells:
#                     cell_text = cell.text.strip()
#                     if cell_text:
#                         text_parts.append(cell_text)

#         # Remove exact consecutive duplicates
#         cleaned_text = []
#         prev = None
#         for line in text_parts:
#             if line != prev:
#                 cleaned_text.append(line)
#             prev = line

#         return "\n".join(cleaned_text)

#     except Exception as e:
#         return f"An error occurred: {e}"

def extract_text_from_docx(file_path):
    try:
        with open(file_path, "rb") as docx_file:
            result = mammoth.extract_raw_text(docx_file)
            text = result.value
            print(text)
            import re
            # text = re.sub(r'[\r\n]+', ' ', text)  # Replace newlines with space
            # text = re.sub(r'\s+', ' ', text)      # Collapse multiple spaces
            text = re.sub(r'(?<! )\n', ' \n', text) # Replaces with newlines with space
            text = re.sub(r'\)(\S)', r') \1', text) # Put space after ")"
            text = re.sub(r'(?<!\s)\(', r' (', text)
            print(text)
            return text.strip()  # plain text
    except Exception as e:
        return f"An error occurred: {e}"


def extract_text_from_txt(file_path):
    """Extracts text from a TXT file."""
    with open(file_path, 'r', encoding='utf-8') as file:
        text = file.read()
    print("-------------------------------")
    print(f"RESUME TEXT: {text}")
    
    return text.strip()

def extract_text(file_path):
    """Extracts text based on file type."""
    _, file_extension = os.path.splitext(file_path)
    file_extension = file_extension.lower()

    if file_extension == ".pdf":
        return extract_text_from_pdf(file_path)
    elif file_extension == ".docx":
        return extract_text_from_docx(file_path)
    elif file_extension == ".txt":
        return extract_text_from_txt(file_path)
    else:
        raise ValueError("Unsupported file format. Please upload a PDF, DOCX, or TXT file.")
